# -*- coding: utf-8 -*-
"""Build the SSVEP FBCCA/Ridge/LRT technical report DOCX."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "SSVEP_FBCCA_Ridge5_LRT算法报告.docx"

BASE_FONT = "Calibri"
EA_FONT = "Microsoft YaHei"
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 105, 117)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TABLE_HEADER = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
LIGHT_FILL = "F8FAFC"
BORDER = "C9D2DD"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_east_asian_font(run, font_name: str = BASE_FONT, ea_font: str = EA_FONT) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    rfonts.set(qn("w:eastAsia"), ea_font)


def set_run(run, *, size: float | None = None, bold: bool | None = None, italic: bool | None = None,
            color: RGBColor | None = None, font: str = BASE_FONT, ea_font: str = EA_FONT) -> None:
    set_east_asian_font(run, font, ea_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6, line: float = 1.25,
                          keep_with_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next


def set_style_font(style, *, size: float, color: RGBColor | None = None, bold: bool | None = None,
                   font: str = BASE_FONT, ea_font: str = EA_FONT) -> None:
    style.font.name = font
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:cs"), font)
    rfonts.set(qn("w:eastAsia"), ea_font)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        set_style_font(style, size=11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_section(section) -> None:
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run(run, size=9, color=MUTED)


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, after=0, line=1.0)
    r = hp.add_run("SSVEP 算法报告")
    set_run(r, size=9, color=MUTED, bold=True)
    r = hp.add_run(" | FBCCA + Ridge5 + LRT")
    set_run(r, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(fp, after=0, line=1.0)
    r = fp.add_run("Page ")
    set_run(r, size=9, color=MUTED)
    add_page_field(fp)


def paragraph_border_bottom(paragraph, color: str = "2E74B5", size: str = "10") -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_p(doc: Document, text: str = "", *, style: str | None = None, size: float | None = None,
          bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None,
          before: float | None = None, after: float | None = None, line: float | None = None,
          align: int | None = None) -> object:
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, italic=italic, color=color)
    if before is not None or after is not None or line is not None:
        set_paragraph_spacing(
            p,
            before=0 if before is None else before,
            after=6 if after is None else after,
            line=1.25 if line is None else line,
        )
    if align is not None:
        p.alignment = align
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_east_asian_font(run)


def cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = tcpr.find(qn("w:tcMar"))
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for tag, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        child = tcmar.find(qn(f"w:{tag}"))
        if child is None:
            child = OxmlElement(f"w:{tag}")
            tcmar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(int(width_dxa)))
    tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths: Sequence[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")
    tbllayout = tblpr.find(qn("w:tblLayout"))
    if tbllayout is None:
        tbllayout = OxmlElement("w:tblLayout")
        tblpr.append(tbllayout)
    tbllayout.set(qn("w:type"), "fixed")

    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{tag}"))
        if edge is None:
            edge = OxmlElement(f"w:{tag}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), BORDER)

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width)))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths[min(idx, len(col_widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.5,
                  color: RGBColor | None = None, monospace: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, font="Consolas" if monospace else BASE_FONT,
            ea_font="Microsoft YaHei Mono" if monospace else EA_FONT)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int], *, font_size: float = 9.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for i, label in enumerate(headers):
        cell_shading(header_cells[i], TABLE_HEADER)
        set_cell_text(header_cells[i], label, bold=True, size=font_size, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=4, line=1.2)
    r = p.add_run(title)
    set_run(r, size=10.5, bold=True, color=INK)
    p = cell.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=1.25)
    r = p.add_run(body)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.text = ""
    for idx, line in enumerate(code.strip("\n").splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(line)
        set_run(r, size=8.8, font="Consolas", ea_font="Microsoft YaHei Mono", color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run(r, size=10.5)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r, size=10.5)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25


def add_metadata_rows(doc: Document, rows: Sequence[tuple[str, str]]) -> None:
    for label, value in rows:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        r = p.add_run(f"{label}: ")
        set_run(r, size=10.5, bold=True, color=INK)
        r = p.add_run(value)
        set_run(r, size=10.5, color=INK)


def add_title_block(doc: Document) -> None:
    add_p(doc, "算法报告", size=10, bold=True, color=BLUE, before=0, after=6, line=1.0)
    p = add_p(
        doc,
        "SSVEP 在线识别算法报告：FBCCA 分数前端 + 5 类 Ridge 分类器 + LRT 多窗口 idle 拒绝门控",
        size=22,
        bold=True,
        color=INK,
        before=0,
        after=4,
        line=1.08,
    )
    p.paragraph_format.keep_with_next = True
    add_p(doc, "基于当前仓库 02_SSVEP 主线源码的实现说明", size=12, color=MUTED, after=14, line=1.15)
    add_metadata_rows(
        doc,
        [
            ("范围", "当前仓库 02_SSVEP 主线；重点覆盖 session no-control 训练、在线解码、idle 拒绝门控"),
            ("生成日期", "2026-05-21"),
            ("核心对象", "FBCCA 分数前端、5 类 Ridge 分类器、LRT multi-window reject gate"),
            ("控制目标", "idle + 8 Hz / 10 Hz / 12 Hz / 15 Hz 四个 SSVEP 指令频率"),
        ]
    )
    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=6, after=14, line=1.0)
    paragraph_border_bottom(rule, color="2E74B5", size="12")


def build() -> None:
    doc = Document()
    configure_section(doc.sections[0])
    configure_styles(doc)
    configure_header_footer(doc)

    add_title_block(doc)
    add_callout(
        doc,
        "一句话结论",
        "当前主线不是“FBCCA 得分最高就输出”的简单解码器，而是把 FBCCA 作为稳定、可解释的分数前端；"
        "再用带 idle 类的 Ridge 线性分类器把每个滑窗映射成 5 类概率；最后用 LRT 证据和多窗口状态机决定是否真正进入 selected 状态。"
        "门控层是压低 no-control/idle 误触发的关键。"
    )

    add_heading(doc, "1. 当前仓库里的实现定位", 1)
    add_p(
        doc,
        "仓库 README 中给出的当前实用主线是：custom session collection -> session_manifest.json + raw_trials.npz "
        "-> fbcca_ridge5 score classifier -> full_reference_bank features -> lrt_multiwindow_reject_gate "
        "-> session no-control calibration -> realtime profile selected by the operator。本文按这条链路解释源码，而不是泛化介绍所有 SSVEP 算法。"
    )
    add_table(
        doc,
        ["模块", "源码位置", "职责"],
        [
            [
                "训练入口/UI",
                "02_SSVEP/apps/realtime_online_ui.py",
                "收集 session no-control 预训练片段，保存数据集 bundle，并在对应模式下调用 Ridge5 + LRT profile 拟合流程。",
            ],
            [
                "离线拟合器",
                "02_SSVEP/ssvep_core/session_no_control_classifier.py",
                "提取 FBCCA 分数特征、训练 5 类 Ridge、拟合 LRT idle 拒绝统计量，并生成 runtime-safe ThresholdProfile。",
            ],
            [
                "分数特征与分类运行时",
                "02_SSVEP/ssvep_core/score_classifier_runtime.py",
                "定义 feature names、分数矩阵到特征的转换、Ridge softmax 预测、LRT 单窗证据计算。",
            ],
            [
                "FBCCA 计算核",
                "02_SSVEP/ssvep_core/compute_kernels.py",
                "构造谐波参考、预处理滑窗、滤波器组、CCA 相关系数与 FBCCA 加权分数。",
            ],
            [
                "在线解码与门控",
                "02_SSVEP/ssvep_core/async_fbcca_idle_standalone.py",
                "加载 profile，创建 FBCCAScoreRidge5ClassifierDecoder，并通过 AsyncDecisionGate 做多窗口 selected/idle 状态管理。",
            ],
        ],
        [1700, 3000, 4660],
    )
    add_callout(
        doc,
        "实现注意",
        "当前检入的 default_profile.json 仍是便携的 baseline FBCCA profile，不是训练好的 fbcca_score_ridge_5class。"
        "Ridge5 + LRT profile 是 session no-control 预训练后生成并由操作者在实时界面选择的 profile；代码也明确要求不要自动覆盖 default_profile.json。"
    )

    add_heading(doc, "2. 总体执行架构", 1)
    add_p(
        doc,
        "这条算法链可以理解成三层：第一层把 EEG 窗口变成可解释的频率匹配分数；第二层把分数形状变成“idle 或某个指令频率”的概率；"
        "第三层用时间上的连续证据过滤掉偶然高分窗口。"
    )
    add_code_block(
        doc,
        """
EEG stream / calibration trials
  -> 2.0 s sliding windows, 0.25 s step
  -> preprocessing: baseline subtraction, optional notch, mean centering
  -> FBCCA scorer over command bank: 8, 10, 12, 15 Hz
  -> FBCCA scorer over full reference bank: 8, 9, 10, 11, 12, 13, 14, 15 Hz
  -> score-shape feature vector
  -> Ridge5 classifier: P(idle), P(8), P(10), P(12), P(15)
  -> probability smoothing + LRT window evidence
  -> AsyncDecisionGate: candidate -> selected -> idle
  -> selected_freq or None
        """
    )
    add_table(
        doc,
        ["阶段", "输入", "输出", "设计目的"],
        [
            ["FBCCA 前端", "多通道 EEG 滑窗", "每个候选频率的相关性分数", "保留 SSVEP 的物理可解释性，避免直接对短校准数据训练复杂模型。"],
            ["分数特征", "4 指令分数 + 8 频全参考分数", "top1、margin、ratio、entropy、non-command margin 等特征", "让分类器看到“像不像一个可靠指令”的分数形状。"],
            ["5 类 Ridge", "标准化后的特征矩阵", "idle/8/10/12/15 五类 softmax 概率", "把 idle 明确建模成一类，而不是只靠阈值事后拒绝。"],
            ["LRT 多窗口门控", "单窗 LRT 证据、概率、分数形状", "selected_freq 或 None", "用 no-control 统计和连续窗口一致性降低误触发。"],
        ],
        [1500, 2250, 2300, 3310],
    )

    add_heading(doc, "3. FBCCA 分数前端原理", 1)
    add_heading(doc, "3.1 SSVEP 参考信号", 2)
    add_p(
        doc,
        "SSVEP 的核心假设是：当用户注视频率为 f 的闪烁刺激时，枕区 EEG 会在 f 及其谐波处产生稳定响应。"
        "源码中的 reference tensor 为每个候选频率构造 sin/cos 谐波基："
    )
    add_code_block(
        doc,
        """
Y_f(t) = [sin(2*pi*1*f*t), cos(2*pi*1*f*t),
          ...,
          sin(2*pi*Nh*f*t), cos(2*pi*Nh*f*t)]
        """
    )
    add_p(
        doc,
        "session no-control classifier 的 scorer 参数为 Nh=5、subband_weight_mode=chen_fixed。也就是说每个频率的参考矩阵包含 1 到 5 次谐波的正弦/余弦列，"
        "构造后会按列减均值，避免 DC 偏置进入 CCA。"
    )

    add_heading(doc, "3.2 预处理与滤波器组", 2)
    add_p(
        doc,
        "每个滑窗进入 FBCCA 前会先做低频基线估计并相减，再按需要做工频陷波，最后按时间轴中心化。随后同一窗口会经过一组带通子带："
        "6-50 Hz、10-50 Hz、14-50 Hz、18-50 Hz、22-50 Hz。子带越靠前通常覆盖越完整的低频 SSVEP 成分，后续子带更强调高频/谐波结构。"
    )
    add_p(
        doc,
        "chen_fixed 权重在源码里按 w_m ∝ (m+1)^(-1.25)+0.25 生成并归一化。FBCCA 因此不是简单平均多个子带，而是让低阶子带权重更大，同时不完全丢弃高阶子带。"
    )

    add_heading(doc, "3.3 CCA 与 FBCCA 分数", 2)
    add_p(
        doc,
        "CCA 会寻找 EEG 子带窗口 X_m 与参考矩阵 Y_f 的两组线性投影，使投影后的相关系数最大。源码中通过协方差白化和 SVD 取最大奇异值，得到每个子带、每个频率的最大 canonical correlation，记为 rho_m,f。"
    )
    add_code_block(
        doc,
        """
rho_m,f = max corr(a^T X_m, b^T Y_f)
score_f = sum_m w_m * rho_m,f^2
predicted_by_plain_fbcca = argmax_f score_f
        """
    )
    add_p(
        doc,
        "在本算法中，argmax_f 不是最终输出；FBCCA 分数只是前端特征来源。这样做的好处是：保留传统 FBCCA 的鲁棒频率匹配能力，同时把 idle、非目标频率泄漏、分数不稳定等问题交给后续分类器和门控层处理。"
    )

    add_heading(doc, "4. 分数特征构造", 1)
    add_p(
        doc,
        "Ridge5 的输入不是原始 EEG，而是 score_classifier_runtime.py 从 FBCCA 分数矩阵构造的特征。对 4 个指令频率，特征包含原始分数和一组描述分数形状的派生量。"
    )
    add_table(
        doc,
        ["特征组", "字段", "含义"],
        [
            ["指令原始分数", "fbcca_score_8, fbcca_score_10, fbcca_score_12, fbcca_score_15", "四个指令频率的 FBCCA 匹配强度。"],
            ["top 分数", "top1_score, top2_score", "当前窗口最强和次强指令分数。"],
            ["分离度", "margin = top1 - top2; ratio = top1 / top2", "目标峰值是否明显压过其他指令。"],
            ["归一化强度", "normalized_top1 = top1 / sum(scores)", "最高分在总能量中的占比。"],
            ["不确定性", "score_entropy", "把指令分数归一化成概率形状后的熵；越高说明越像多个频率混在一起。"],
            ["全参考库", "top_command_to_top_all_ratio, nearest_noncommand_margin, all_bank_entropy 等", "比较 4 个指令频率和 8 个参考频率，帮助识别非指令频率或 idle 干扰。"],
        ],
        [1500, 3000, 4860],
        font_size=9.2,
    )
    add_p(
        doc,
        "full_reference_bank 是这版架构的重要增强：训练器会同时跑一个 4 频 command decoder 和一个 8 频 full-bank decoder。"
        "如果 9/11/13/14 Hz 这类非指令参考在窗口里也很强，nearest_noncommand_margin 和 all_bank_entropy 会把这种风险暴露给 Ridge/LRT。"
    )

    add_heading(doc, "5. 5 类 Ridge 分类器", 1)
    add_heading(doc, "5.1 为什么是 5 类", 2)
    add_p(
        doc,
        "分类标签顺序固定为 idle + 4 个指令频率，即 idle、8、10、12、15。这样 idle 不再只是“所有频率都不够强”的残差情况，而是有自己的校准样本、均值方差和分类边界。"
        "对于需要控制机器人或 UI 的在线系统，这比纯阈值拒绝更安全，因为模型学习的是“真实 no-control 窗口的分数形状”。"
    )

    add_heading(doc, "5.2 训练公式", 2)
    add_p(
        doc,
        "训练只使用 calibration trial。每个 trial 被切成多个 2 秒滑窗；每个窗口都有一行 FBCCA 特征，并继承该 trial 的标签。"
        "为了避免某一类窗口数量多而主导拟合，源码按类别窗口数的倒数设置 sample weight，然后做闭式 Ridge 求解。"
    )
    add_code_block(
        doc,
        """
X        : calibration feature matrix
mu, sigma: feature-wise mean and std
Z        = (X - mu) / sigma
D        = [1, Z]                         # intercept + standardized features
T        = one_hot(labels)                 # idle, 8, 10, 12, 15
S        = diagonal sample-weight matrix   # inverse class frequency, normalized
R        = identity, with R[0,0] = 0        # do not regularize intercept
W        = (D^T S D + lambda R)^(-1) D^T S T
lambda   = 0.3
        """
    )
    add_p(
        doc,
        "运行时同样先用保存的 mu/sigma 标准化特征，再计算 logits = [1, Z]W。源码随后对 logits 做 softmax，得到五类概率。"
        "Ridge 本质上是一个轻量线性分类头：校准样本少时比深度模型更稳定，也更容易审计。"
    )

    add_heading(doc, "5.3 在线概率与命令置信度", 2)
    add_p(
        doc,
        "FBCCAScoreRidge5ClassifierDecoder 输出 classifier_probs、classifier_pred_label、classifier_pred_freq 和 classifier_command_confidence。"
        "其中 command_confidence = 1 - P(idle)。源码还会对概率做 smoothing_windows=3 的滑动平均，减少单个窗口抖动导致的门控误判。"
    )

    add_heading(doc, "6. LRT 多窗口 idle 拒绝门控", 1)
    add_heading(doc, "6.1 单窗口 LRT 证据", 2)
    add_p(
        doc,
        "LRT 是 log likelihood ratio，即“这个窗口更像控制状态还是更像 idle 状态”的对数似然差。训练时会从校准数据中取两组样本："
        "idle trial 的所有窗口作为 idle 分布；非 idle trial 中 Ridge 平滑后预测正确的窗口作为 control 分布。"
    )
    add_p(
        doc,
        "用于 LRT 的特征不是全部特征，而是 top1_score、margin、ratio、normalized_top1、score_entropy、top_command_to_top_all_ratio、nearest_noncommand_margin、all_bank_entropy。"
        "这些特征共同描述“强不强、尖不尖、是不是被非指令频率抢走”。"
    )
    add_code_block(
        doc,
        """
z_control = (x - mean_control) / std_control
z_idle    = (x - mean_idle)    / std_idle

log_control = -0.5 * sum(z_control^2) - sum(log(std_control))
log_idle    = -0.5 * sum(z_idle^2)    - sum(log(std_idle))
LRT(x)      = log_control - log_idle
        """
    )
    add_p(
        doc,
        "当 LRT(x) 为正且足够大，窗口更接近控制分布；当它偏低或为负，窗口更像 idle/no-control。这里使用的是对角高斯近似，计算量小，适合实时门控。"
    )

    add_heading(doc, "6.2 阈值如何来", 2)
    add_p(
        doc,
        "训练器把 idle 窗口的 LRT 第 95 百分位作为拒绝线候选，又用 control 窗口的第 50 百分位进行上限约束，最后取不小于 0 的值。直觉上："
    )
    add_bullets(
        doc,
        [
            "阈值高于大多数 idle 窗口，可以压低 idle 误触发。",
            "阈值不超过典型 control 窗口太多，避免把正常控制窗口全部挡掉。",
            "阈值下限为 0，要求窗口至少比 idle 更像 control。",
        ]
    )
    add_p(
        doc,
        "源码中 lrt_enter_th 当前设为 0.0，因此选中主要依赖单窗 LRT 过线后的同标签连续窗口数；同时保留 evidence_by_label 的累积与衰减机制，便于后续变体提高进入门槛。"
    )

    add_heading(doc, "6.3 多窗口状态机", 2)
    add_p(
        doc,
        "AsyncDecisionGate 不会因为一个窗口分类成 10 Hz 就立刻输出 10 Hz。它维护 idle、candidate、selected 三种主要状态，以及每个标签的累计证据、同标签 streak、gap 和 exit 计数。"
    )
    add_table(
        doc,
        ["检查项", "源码逻辑", "作用"],
        [
            ["非 idle 预测", "classifier_pred_freq 必须存在", "如果 Ridge 当前认为是 idle，直接拒绝。"],
            ["命令置信度", "1 - P(idle) 需超过 command_confidence_th", "避免 P(idle) 仍很高时进入控制。"],
            ["LRT 窗口证据", "lrt_window_evidence >= max(lrt_window_th, lrt_window_floor_th)", "确认窗口更像控制分布。"],
            ["分数形状", "margin/ratio/entropy 等 score-shape guard", "防止多频混叠或峰值不尖锐的窗口进入。"],
            ["连续窗口", "同一 pred_label streak >= min_enter_windows，默认 2", "要求时间上稳定一致。"],
            ["退出", "selected 后当前窗口不再通过检查，连续 min_exit_windows 后回 idle", "让用户松开注视或信号变差时释放控制。"],
        ],
        [1550, 3350, 4460],
        font_size=9.2,
    )

    add_heading(doc, "7. 训练与标定流程", 1)
    add_p(
        doc,
        "session no-control 预训练模式会采集四个指令频率和 idle/no-control 片段。默认控制频率为 8、10、12、15 Hz；训练窗口为 2.0 s，步长 0.25 s。"
        "UI 会先保存 session_manifest.json 与 raw_trials.npz 之类的数据集 bundle，再拟合 profile。"
    )
    add_numbered(
        doc,
        [
            "采集 command trial 与 idle trial，记录采样率、通道、频率和 trial 元数据。",
            "每个 trial 切成固定长度滑窗，分别用 command bank 和 full reference bank 计算 FBCCA 分数矩阵。",
            "把分数矩阵转换为特征矩阵，检查 idle 与每个指令类都有校准样本。",
            "用平衡 sample weights 训练 Ridge5，保存 feature_mean、feature_std、weights、labels 和 l2。",
            "用 Ridge 平滑预测挑出可靠 control 窗口，和 idle 窗口一起拟合 LRT 的均值、方差和阈值。",
            "生成 model_name=fbcca_score_ridge_5class、gate_policy=lrt_multiwindow_reject_gate 的 ThresholdProfile。",
            "保存到 session no-control classifier profile 路径和历史路径，由实时界面操作者选择使用。",
        ]
    )

    add_heading(doc, "8. 在线执行流程", 1)
    add_p(
        doc,
        "实时阶段会先加载 profile，再按 profile 的 model_name 创建解码器。如果 profile 是 fbcca_score_ridge_5class，运行时要求 model_params.state 非空；"
        "否则会拒绝启动，因为 Ridge 权重、标准化统计量和 LRT 参数都必须来自当前被试/当前 session 的训练结果。"
    )
    add_code_block(
        doc,
        """
profile = load_profile(...)
decoder = load_decoder_from_profile(profile, sampling_rate=fs, compute_backend=cpu/cuda)
gate    = AsyncDecisionGate.from_profile(profile)

for each incoming EEG window:
    analysis = decoder.analyze_window(window)
    # analysis contains FBCCA scores, Ridge probabilities, pred label, LRT evidence
    decision = gate.update(analysis)
    emit decision["selected_freq"] if decision["state"] == "selected" else None
        """
    )
    add_p(
        doc,
        "因此，在线输出 selected_freq 是门控后的控制事件，而 pred_freq 只是当前窗口的候选预测。下游控制模块应优先消费 selected_freq，"
        "因为它已经经过 idle 拒绝、连续窗口确认和退出逻辑。"
    )

    add_heading(doc, "9. 架构特点、收益与限制", 1)
    add_heading(doc, "9.1 主要收益", 2)
    add_bullets(
        doc,
        [
            "可解释性强：FBCCA 分数、Ridge 概率和 LRT 证据都能被日志记录和回放分析。",
            "短校准友好：Ridge 是闭式线性模型，适合每个 session 快速拟合。",
            "idle 显式建模：idle 是第 5 类，并且 LRT 还使用 no-control 分布二次把关。",
            "抗非指令频率：full reference bank 会暴露非指令参考频率抢占最高分的情况。",
            "实时稳定：多窗口 streak、evidence decay 和 exit 计数减少瞬时误判。",
        ]
    )
    add_heading(doc, "9.2 使用限制", 2)
    add_bullets(
        doc,
        [
            "必须有足够的 idle 和四个指令类校准窗口；缺任何类都会导致训练失败或 profile 不可靠。",
            "2 秒窗口决定了最低感知延迟；0.25 秒步长只提高更新频率，不会让第一个可靠窗口早于 2 秒。",
            "LRT 使用对角高斯近似，假设选中特征近似独立；它是工程上快速稳定的门控，不是完整概率生成模型。",
            "门控偏保守时会降低 recall；阈值偏松时会提高 idle false positive，需要结合实时任务风险调参。",
            "当前仓库的默认 baseline profile 可用于 fallback，但不能代表训练后的 Ridge5 + LRT 主线性能。",
        ]
    )

    doc.add_page_break()
    add_heading(doc, "附录 A. 关键参数表", 1)
    add_table(
        doc,
        ["参数", "当前实现值", "说明"],
        [
            ["控制频率", "8, 10, 12, 15 Hz", "四个可输出的 SSVEP 指令频率。"],
            ["第 5 类", "idle", "no-control/不注视指令时的显式类别。"],
            ["full reference bank", "8, 9, 10, 11, 12, 13, 14, 15 Hz", "用于构造非指令频率相关特征。"],
            ["win_sec / step_sec", "2.0 s / 0.25 s", "训练和在线滑窗设置。"],
            ["Nh", "5", "Ridge5 主线 scorer 使用 5 次谐波。"],
            ["子带", "6-50, 10-50, 14-50, 18-50, 22-50 Hz", "FBCCA 默认五子带。"],
            ["子带权重", "chen_fixed: a=1.25, b=0.25", "按 (m+1)^(-a)+b 生成并归一化。"],
            ["Ridge L2", "0.3", "权重正则；截距不正则。"],
            ["smoothing_windows", "3", "Ridge 概率滑动平均窗口数。"],
            ["min_enter_windows", "2", "默认需要两个连续通过窗口才能 selected。"],
            ["min_exit_windows", "1", "默认一个失败窗口即可退出 selected。"],
            ["model_name", "fbcca_score_ridge_5class", "实时加载该解码器的模型名。"],
            ["gate_policy", "lrt_multiwindow_reject_gate", "LRT + 多窗口 idle 拒绝门控。"],
            ["schema_version", "session_no_control_fbcca_ridge5_v1", "session no-control classifier state 版本。"],
        ],
        [2200, 2600, 4560],
        font_size=9.2,
    )

    add_heading(doc, "附录 B. 训练伪代码", 1)
    add_code_block(
        doc,
        """
def fit_session_profile(trials):
    scored_trials = []
    for trial in trials:
        windows = extract_window_batch(trial.eeg, win=2.0s, step=0.25s)
        command_scores = FBCCA(freqs=[8,10,12,15], Nh=5).score(windows)
        all_scores = FBCCA(freqs=[8,9,10,11,12,13,14,15], Nh=5).score(windows)
        features = score_matrices_to_features(command_scores, all_scores)
        scored_trials.append((trial.label, features))

    ridge_state = fit_balanced_ridge(scored_trials, labels=[idle,8,10,12,15], l2=0.3)
    lrt_state = fit_lrt_gate_state(ridge_state, scored_trials)
    profile = ThresholdProfile(
        model_name="fbcca_score_ridge_5class",
        model_params={"state": ridge_state + lrt_state},
        gate_policy="lrt_multiwindow_reject_gate",
    )
    return profile
        """
    )

    add_heading(doc, "附录 C. 在线伪代码", 1)
    add_code_block(
        doc,
        """
def realtime_step(window, decoder, gate):
    features = decoder.analyze_window(window)
    # features include:
    #   scores, classifier_probs, classifier_pred_freq,
    #   classifier_command_confidence, lrt_window_evidence
    decision = gate.update(features)
    if decision["state"] == "selected":
        return decision["selected_freq"]
    return None
        """
    )

    add_heading(doc, "附录 D. 源码索引", 1)
    add_table(
        doc,
        ["文件", "关键符号", "阅读重点"],
        [
            ["02_SSVEP/README.md", "Current Mainline", "确认当前实用路径和不要覆盖 default_profile 的约束。"],
            ["02_SSVEP/ssvep_core/session_no_control_classifier.py", "fit_session_no_control_fbcca_ridge_profile", "从校准片段到 runtime profile 的完整拟合流程。"],
            ["02_SSVEP/ssvep_core/session_no_control_classifier.py", "_fit_balanced_ridge_classifier", "Ridge5 的标准化、类别均衡、闭式求解。"],
            ["02_SSVEP/ssvep_core/session_no_control_classifier.py", "_fit_lrt_gate_state", "LRT 特征选择、control/idle 分布、阈值选择。"],
            ["02_SSVEP/ssvep_core/score_classifier_runtime.py", "score_matrices_to_features", "FBCCA 分数转特征矩阵。"],
            ["02_SSVEP/ssvep_core/score_classifier_runtime.py", "ridge5_predict_windows_from_state", "Ridge logits + softmax 概率。"],
            ["02_SSVEP/ssvep_core/score_classifier_runtime.py", "lrt_window_evidence_from_state", "对角高斯 LRT 证据计算。"],
            ["02_SSVEP/ssvep_core/compute_kernels.py", "fbcca_scores_batch", "预处理、子带滤波、CCA、加权求分。"],
            ["02_SSVEP/ssvep_core/async_fbcca_idle_standalone.py", "FBCCAScoreRidge5ClassifierDecoder", "在线窗口分析，输出概率、特征和 LRT 证据。"],
            ["02_SSVEP/ssvep_core/async_fbcca_idle_standalone.py", "AsyncDecisionGate.update", "候选、进入、保持、退出 selected 状态的状态机。"],
            ["02_SSVEP/apps/realtime_online_ui.py", "Pretrain worker / profile paths", "UI 如何采集 session、保存 profile，并在实时运行中加载 decoder 和 gate。"],
        ],
        [2900, 2550, 3910],
        font_size=8.8,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    build()
